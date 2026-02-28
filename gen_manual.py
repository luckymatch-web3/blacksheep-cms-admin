"""Generate CMS Admin User Manual as DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Microsoft YaHei'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()


def add_tip(text):
    """Add a tip/note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"💡 提示：{text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x00)
    run.font.italic = True


def add_warning(text):
    """Add a warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ 注意：{text}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    run.font.italic = True


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BlackSheep CMS")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("内容管理系统 使用手册")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x40, 0x9E, 0xFF)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("版本 1.0  |  2026 年 2 月")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading("目录", level=1)
toc_items = [
    "1. 系统概述",
    "2. 登录系统",
    "3. 仪表盘（Dashboard）",
    "4. 文章管理",
    "    4.1 文章列表",
    "    4.2 新建文章",
    "    4.3 编辑文章",
    "    4.4 文章状态说明",
    "5. 轮播管理（Banner）",
    "6. 文章审核",
    "7. 帖子审核",
    "8. 回收站",
    "9. 提现管理",
    "10. 常见问题",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3

doc.add_page_break()

# ============================================================
# 1. 系统概述
# ============================================================
doc.add_heading("1. 系统概述", level=1)

doc.add_paragraph(
    "BlackSheep CMS 是 BlackSheep Media 的内容管理后台，用于管理网站首页文章、"
    "轮播 Banner、论坛帖子审核以及用户提现审批等核心业务。"
)

doc.add_heading("访问地址", level=3)
doc.add_paragraph("https://blacksheepmedia.xyz/cms/")

doc.add_heading("系统要求", level=3)
doc.add_paragraph("推荐使用 Chrome、Edge 或 Firefox 最新版本浏览器。系统采用深色主题界面。")

doc.add_heading("管理员账号", level=3)
add_table(
    ["项目", "值"],
    [
        ["用户名", "testlocal"],
        ["密码", "admin123"],
        ["所需权限", "ROLE_ADMIN"],
    ],
)

doc.add_heading("功能模块概览", level=3)
add_table(
    ["模块", "功能说明"],
    [
        ["仪表盘", "数据统计总览、快捷操作入口"],
        ["文章列表", "文章的增删改查、状态管理"],
        ["新建/编辑文章", "富文本编辑器，支持封面图、分类、标签、Banner 设置"],
        ["轮播管理", "首页 Banner 轮播的添加、排序、移除"],
        ["文章审核", "审核待发布文章，通过或驳回"],
        ["帖子审核", "审核用户论坛帖子，通过后公开显示"],
        ["回收站", "恢复或永久删除已删除的文章"],
        ["提现管理", "审核用户 USDT 提现申请，打款或驳回"],
    ],
)

doc.add_page_break()

# ============================================================
# 2. 登录系统
# ============================================================
doc.add_heading("2. 登录系统", level=1)

doc.add_paragraph(
    "打开系统地址后，如未登录会自动跳转到登录页面。"
)

doc.add_heading("登录步骤", level=3)
steps = [
    "在浏览器中打开 https://blacksheepmedia.xyz/cms/",
    "输入管理员用户名和密码",
    "点击「登录」按钮或按 Enter 键",
    "登录成功后自动跳转到仪表盘",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

add_tip("登录页面已预填默认用户名 testlocal，只需输入密码即可。")

doc.add_heading("登录失败处理", level=3)
add_table(
    ["错误情况", "处理方式"],
    [
        ["用户名或密码错误", "页面下方显示红色错误提示，请检查后重试"],
        ["无管理员权限", '提示"该账号没有管理员权限"，请联系超级管理员'],
        ["网络错误", '提示"网络错误，请稍后重试"，检查网络连接'],
    ],
)

doc.add_heading("退出登录", level=3)
doc.add_paragraph(
    "点击页面右上角用户头像旁的「退出」按钮即可退出登录，系统将跳转回登录页面。"
)

add_warning('登录状态过期（Token 失效）时，系统会自动跳转到登录页面，提示"登录已过期"。')

doc.add_page_break()

# ============================================================
# 3. 仪表盘
# ============================================================
doc.add_heading("3. 仪表盘（Dashboard）", level=1)

doc.add_paragraph(
    "仪表盘是登录后的默认首页，提供内容数据的全局概览和常用操作的快捷入口。"
)

doc.add_heading("数据统计卡片", level=3)
doc.add_paragraph("页面顶部显示 6 个统计卡片：")
add_table(
    ["卡片", "说明", "颜色"],
    [
        ["全部文章", "系统中所有文章总数", "蓝色"],
        ["已发布", '状态为"已发布"的文章数量', "绿色"],
        ["草稿", '状态为"草稿"的文章数量', "灰色"],
        ["待审核", '状态为"待审核"的文章数量', "黄色"],
        ["轮播", "当前设置为 Banner 的文章数量", "黄色"],
        ["回收站", "已删除（在回收站中）的文章数量", "灰色"],
    ],
)

doc.add_heading("最近文章", level=3)
doc.add_paragraph(
    "左侧面板显示最近 5 篇文章，包含标题、分类、日期和状态标签。"
    "点击任意文章可直接跳转到编辑页面。"
)

doc.add_heading("快捷操作", level=3)
doc.add_paragraph("右侧面板提供 4 个快捷按钮：")
add_table(
    ["按钮", "功能"],
    [
        ["新建文章", "跳转到文章创建页面"],
        ["轮播管理", "跳转到 Banner 管理页面"],
        ["内容审核", "跳转到文章审核页面"],
        ["查看前台", "在新标签页打开网站首页"],
    ],
)

doc.add_page_break()

# ============================================================
# 4. 文章管理
# ============================================================
doc.add_heading("4. 文章管理", level=1)

# 4.1
doc.add_heading("4.1 文章列表", level=2)
doc.add_paragraph(
    "文章列表页面展示所有文章，支持按关键词搜索和按状态筛选。"
    "在侧边栏点击「文章列表」即可进入。"
)

doc.add_heading("搜索与筛选", level=3)
doc.add_paragraph(
    "页面顶部提供搜索栏：\n"
    "• 关键词输入框：输入文章标题关键词，按 Enter 或点击「搜索」按钮\n"
    "• 状态下拉框：可选择 全部 / 草稿 / 审核中 / 已发布 / 已驳回\n"
    "• 重置按钮：清除所有筛选条件"
)

doc.add_heading("文章列表表格", level=3)
doc.add_paragraph("表格包含以下列：")
add_table(
    ["列名", "说明"],
    [
        ["ID", "文章编号，格式为 #123"],
        ["文章", "封面缩略图 + 文章标题 + 副标题"],
        ["分类", "文章所属分类"],
        ["状态", "当前状态标签（颜色区分）"],
        ["轮播", "如果是 Banner 则显示星标和排序号"],
        ["浏览", "文章浏览次数"],
        ["日期", "创建日期"],
        ["操作", "可执行的操作按钮（因状态而异）"],
    ],
)

doc.add_heading("操作按钮说明", level=3)
doc.add_paragraph("不同状态下可用的操作按钮不同：")
add_table(
    ["文章状态", "可用操作"],
    [
        ["草稿", "编辑、提交审核、直接发布、删除"],
        ["待审核", "编辑、通过、驳回"],
        ["已发布", "编辑、下架、删除"],
        ["已驳回", "编辑、重新提交审核、删除"],
    ],
)

add_tip("删除操作是「软删除」，文章会移入回收站，可以恢复。")

# 4.2
doc.add_heading("4.2 新建文章", level=2)
doc.add_paragraph(
    "点击文章列表页的「新建文章」按钮，或侧边栏的「新建文章」菜单项，进入文章创建页面。"
)

doc.add_heading("页面布局", level=3)
doc.add_paragraph("页面分为左右两栏：")
doc.add_paragraph(
    "左侧 — 文章内容区：\n"
    "• 标题（必填）：文章主标题\n"
    "• 副标题（选填）：文章副标题\n"
    "• 摘要（选填）：在列表页显示的简要描述\n"
    "• 正文内容：富文本编辑器"
)
doc.add_paragraph(
    "右侧 — 文章属性区：\n"
    "• 封面图片：支持拖拽上传或输入图片 URL，建议尺寸 800×450\n"
    "• 作者：文章作者名称\n"
    "• 分类：输入分类名，支持自动补全已有分类\n"
    "• 标签：多个标签用英文逗号分隔\n"
    "• Banner 设置：勾选后将文章设为首页轮播，可设置排序值"
)

doc.add_heading("富文本编辑器", level=3)
doc.add_paragraph("编辑器工具栏支持以下格式：")
add_table(
    ["功能", "说明"],
    [
        ["标题", "支持 H1、H2、H3 三级标题"],
        ["文字格式", "加粗、斜体、下划线、删除线"],
        ["颜色", "文字颜色、背景色"],
        ["列表", "有序列表、无序列表"],
        ["引用", "块引用"],
        ["代码", "代码块"],
        ["链接", "插入超链接"],
        ["图片", "插入图片"],
        ["清除格式", "清除所有格式"],
    ],
)

doc.add_heading("保存方式", level=3)
doc.add_paragraph("页面顶部提供三个保存按钮：")
add_table(
    ["按钮", "说明"],
    [
        ["保存草稿", "保存为草稿状态，不会公开显示"],
        ["提交审核", "保存并提交审核，等待管理员审核通过后发布"],
        ["直接发布", "保存并立即发布，文章直接公开可见"],
    ],
)

add_tip("标题是唯一的必填项，其他字段均为选填。")

# 4.3
doc.add_heading("4.3 编辑文章", level=2)
doc.add_paragraph(
    "在文章列表中点击「编辑」按钮即可进入编辑页面，界面与新建文章相同，"
    "但会预填已有的文章内容。"
)
doc.add_paragraph(
    "如果文章曾被驳回，编辑页面顶部会显示红色提示框，包含驳回原因和审核人信息，"
    "方便作者了解需要修改的内容。"
)

# 4.4
doc.add_heading("4.4 文章状态说明", level=2)
doc.add_paragraph("文章在系统中有 4 种状态，遵循以下流转规则：")

add_table(
    ["状态", "颜色标签", "说明"],
    [
        ["草稿 (DRAFT)", "灰色", "新建或恢复的文章默认状态，未公开"],
        ["待审核 (PENDING_REVIEW)", "黄色", "已提交审核，等待管理员处理"],
        ["已发布 (PUBLISHED)", "绿色", "审核通过或直接发布，前台可见"],
        ["已驳回 (REJECTED)", "红色", "审核未通过，需修改后重新提交"],
    ],
)

doc.add_heading("状态流转图", level=3)
doc.add_paragraph(
    "草稿 ──→ 提交审核 ──→ 待审核 ──→ 通过 ──→ 已发布\n"
    "                                  └──→ 驳回 ──→ 已驳回 ──→ 重新提交 ──→ 待审核\n"
    "草稿 ──→ 直接发布 ──→ 已发布\n"
    "已发布 ──→ 下架 ──→ 草稿\n"
    "任何状态 ──→ 删除 ──→ 回收站 ──→ 恢复 ──→ 草稿\n"
    "                              └──→ 永久删除 ──→ 彻底删除"
)

doc.add_page_break()

# ============================================================
# 5. 轮播管理
# ============================================================
doc.add_heading("5. 轮播管理（Banner）", level=1)

doc.add_paragraph(
    "轮播管理用于设置首页的 Banner 轮播图。最多支持 5 个 Banner，只有已发布的文章才能设为 Banner。"
)

doc.add_heading("当前轮播列表", level=3)
doc.add_paragraph(
    "页面上方显示当前所有 Banner，每个 Banner 条目包含：\n"
    "• 排序编号（1-5）\n"
    "• 文章封面缩略图\n"
    "• 文章标题、分类、排序值、浏览量\n"
    "• 操作按钮：上移、下移、编辑、移除"
)

doc.add_heading("操作说明", level=3)
add_table(
    ["操作", "说明"],
    [
        ["上移 / 下移", "调整 Banner 的显示顺序"],
        ["编辑", "跳转到文章编辑页面"],
        ["移除", "取消该文章的 Banner 状态（文章本身不受影响）"],
    ],
)

doc.add_heading("添加 Banner", level=3)
doc.add_paragraph(
    "页面下方提供添加功能：\n"
    "1. 在下拉框中搜索并选择一篇已发布的文章\n"
    "2. 点击「添加」按钮\n"
    "3. 文章将自动添加到轮播列表末尾"
)

add_warning("最多只能添加 5 个 Banner。如需添加新的，请先移除一个现有的 Banner。")

doc.add_page_break()

# ============================================================
# 6. 文章审核
# ============================================================
doc.add_heading("6. 文章审核", level=1)

doc.add_paragraph(
    "文章审核页面只显示状态为「待审核」的文章。当有文章提交审核时，"
    "侧边栏的「文章审核」菜单项会显示红色数字徽标提示。"
)

doc.add_heading("审核操作", level=3)
add_table(
    ["操作", "按钮颜色", "说明"],
    [
        ["查看", "蓝色", "跳转到文章编辑页面查看完整内容"],
        ["通过", "绿色", '审核通过，文章自动变为"已发布"状态，可输入审核备注（选填）'],
        ["驳回", "红色", '审核驳回，文章变为"已驳回"状态，必须填写驳回原因'],
    ],
)

doc.add_heading("审核流程", level=3)
steps = [
    "在审核列表中点击「查看」查看文章完整内容",
    "确认文章内容无误后，点击「通过」并可选填审核备注",
    "如果文章内容不合规或需要修改，点击「驳回」并填写驳回原因",
    "审核完成后，文章状态自动更新，侧边栏徽标数字减 1",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

add_tip("驳回原因会显示在文章编辑页面顶部，便于作者了解修改要求。")

doc.add_page_break()

# ============================================================
# 7. 帖子审核
# ============================================================
doc.add_heading("7. 帖子审核", level=1)

doc.add_paragraph(
    "帖子审核用于管理用户在论坛发布的帖子。用户发布的帖子默认为「待审核」状态，"
    "审核通过后才会在前台公开显示。"
)

doc.add_heading("统计概览", level=3)
doc.add_paragraph(
    "页面顶部显示 4 个统计卡片，同时作为状态筛选按钮：\n"
    "• 待审核（黄色）：点击筛选待审核帖子\n"
    "• 已发布（绿色）：点击筛选已发布帖子\n"
    "• 已驳回（红色）：点击筛选已驳回帖子\n"
    "• 全部：显示所有帖子"
)

doc.add_heading("帖子列表表格", level=3)
add_table(
    ["列名", "说明"],
    [
        ["帖子", "帖子标题或内容摘要"],
        ["分类", "帖子分类（投诉、举报诈骗、正面评价、经验分享等）"],
        ["交易所", "关联的交易所名称和 Logo"],
        ["作者", "发帖用户昵称"],
        ["状态", "待审核 / 已发布 / 已驳回"],
        ["图片", "帖子包含的图片数量"],
        ["发布时间", "帖子创建时间"],
        ["操作", "通过 / 驳回按钮"],
    ],
)

doc.add_heading("审核操作", level=3)
add_table(
    ["帖子状态", "可用操作"],
    [
        ["待审核", "通过（确认对话框）、驳回（需填写原因）"],
        ["已发布", "驳回（撤回已发布的帖子）"],
        ["已驳回", "通过（恢复帖子）"],
    ],
)

add_warning("驳回帖子时必须填写驳回原因，不能为空。")

doc.add_page_break()

# ============================================================
# 8. 回收站
# ============================================================
doc.add_heading("8. 回收站", level=1)

doc.add_paragraph(
    "回收站保存所有被软删除的文章。文章被删除后不会立即消失，"
    "而是进入回收站，管理员可以选择恢复或永久删除。"
)

doc.add_heading("回收站列表", level=3)
doc.add_paragraph(
    "列表中的文章以半透明样式和删除线标题显示，表明已被删除。每条记录显示：\n"
    "• 文章 ID 和标题\n"
    "• 分类\n"
    "• 删除时间\n"
    "• 删除人"
)

doc.add_heading("操作说明", level=3)
add_table(
    ["操作", "按钮颜色", "说明"],
    [
        ["恢复", "绿色", '将文章从回收站恢复，状态变为"草稿"'],
        ["永久删除", "红色", "彻底删除文章，不可恢复。会弹出二次确认对话框"],
    ],
)

add_warning("永久删除操作不可撤销！请确认后再操作。")

doc.add_page_break()

# ============================================================
# 9. 提现管理
# ============================================================
doc.add_heading("9. 提现管理", level=1)

doc.add_paragraph(
    "提现管理用于审核用户的 USDT 提现申请。管理员需人工审核，"
    "确认无误后进行链上打款，并将交易哈希（txHash）回填到系统中。"
)

doc.add_heading("状态筛选", level=3)
doc.add_paragraph(
    "页面顶部提供标签式筛选：\n"
    "• 待处理（黄色）：等待管理员处理的申请\n"
    "• 已完成（绿色）：已打款的申请\n"
    "• 已驳回（红色）：被驳回的申请\n"
    "• 全部：显示所有申请"
)

doc.add_heading("提现记录表格", level=3)
add_table(
    ["列名", "说明"],
    [
        ["提现金额", "申请的 USDT 金额"],
        ["收款地址", "用户的钱包地址（点击复制图标可复制完整地址）"],
        ["用户 ID", "申请用户的 ID"],
        ["状态", "待处理 / 已完成 / 已驳回"],
        ["交易哈希", "链上打款的 txHash（点击复制图标可复制）"],
        ["备注", "申请备注信息"],
        ["申请时间", "提现申请的创建时间"],
        ["操作", "打款 / 驳回按钮"],
    ],
)

doc.add_heading("操作流程", level=3)

doc.add_paragraph("打款流程：", style='List Bullet')
steps = [
    "确认用户提现金额和钱包地址正确",
    "在链上完成 USDT 转账",
    "点击「打款」按钮",
    "在弹出框中输入链上交易哈希（txHash）",
    '确认后申请状态变为"已完成"',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"    {i}. {step}")

doc.add_paragraph("驳回流程：", style='List Bullet')
steps = [
    "点击「驳回」按钮",
    "在弹出框中填写驳回原因（如：钱包地址有误、金额不符等）",
    '确认后申请状态变为"已驳回"',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"    {i}. {step}")

add_tip("点击钱包地址或交易哈希旁的复制图标，可快速复制完整内容。")

doc.add_page_break()

# ============================================================
# 10. 常见问题
# ============================================================
doc.add_heading("10. 常见问题", level=1)

qa = [
    (
        '登录后提示"该账号没有管理员权限"',
        "只有拥有 ROLE_ADMIN 权限的账号才能登录 CMS 后台。请使用管理员账号 testlocal 登录，"
        "或联系系统管理员为你的账号添加管理员权限。"
    ),
    (
        "页面突然跳转到登录页面",
        "登录 Token 有有效期，过期后系统会自动跳转到登录页。重新登录即可，"
        "之前未保存的内容可能会丢失，建议编辑文章时及时保存草稿。"
    ),
    (
        "上传图片失败",
        "请检查：1) 图片格式是否为 JPG/PNG/GIF/WebP；2) 图片大小是否超过限制；"
        "3) 网络连接是否正常。也可以使用图片 URL 方式直接输入在线图片地址。"
    ),
    (
        "文章发布后前台看不到",
        '请确认：1) 文章状态确实是"已发布"；2) 刷新前台页面（浏览器可能有缓存）；'
        "3) 如果是 Banner 文章，检查轮播管理中是否正确设置。"
    ),
    (
        "如何批量管理文章？",
        "目前系统暂不支持批量操作，需要逐篇处理。"
    ),
    (
        "驳回文章后作者如何知道？",
        "作者在编辑文章时会在页面顶部看到红色的驳回提示框，"
        "其中包含驳回原因和审核人信息。"
    ),
    (
        "Banner 最多可以设置几个？",
        "最多 5 个。如需添加新的 Banner，请先移除一个现有的。"
    ),
    (
        "金色财经等数据源无法访问？",
        "部分国内数据源需要国内服务器环境才能正常访问，"
        "本地开发环境可能因网络原因无法连接，在阿里云服务器上运行正常。"
    ),
]

for q, a in qa:
    p = doc.add_paragraph()
    run = p.add_run(f"Q: {q}")
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(f"A: {a}")
    doc.add_paragraph()  # spacer

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— BlackSheep Media © 2026 —")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = r"E:\blacksheep-cms-admin\BlackSheep_CMS_使用手册.docx"
doc.save(output_path)
print(f"Manual saved to: {output_path}")
